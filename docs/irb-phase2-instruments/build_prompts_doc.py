#!/usr/bin/env python3
"""Regenerate the IRB "AI Model System Prompts (Verbatim)" docx + html.

Verbatim prompt blocks are extracted from the TypeScript sources at build time
so the document cannot drift from the deployed code.
"""
import html as html_mod
import re
import sys

REPO = "/Users/nathanblatter/dev/Ai-therapist"
OUT_DOCX = f"{REPO}/docs/irb-phase2-instruments/AI_Model_System_Prompts_Verbatim_09.04.2026.docx"
OUT_HTML = f"{REPO}/docs/irb-phase2-instruments/prompts.html"


def read(path):
    with open(f"{REPO}/{path}", encoding="utf-8") as f:
        return f.read()


def decode_js(s):
    """Decode escape sequences inside a JS template/string literal."""
    return (s.replace("\\n", "\n").replace("\\'", "'").replace('\\"', '"')
             .replace("\\`", "`").replace("\\\\", "\\"))


def extract_backtick(src, anchor):
    """Return the decoded contents of the backtick literal following `anchor`."""
    i = src.index(anchor)
    start = src.index("`", i) + 1
    end = src.index("`", start)
    return decode_js(src[start:end])


sh = read("src/server/utils/sessionHelpers.ts")
cd = read("src/server/services/crisisDetection.service.ts")
red = read("src/server/services/redaction.service.ts")
sn = read("src/server/services/sessionName.service.ts")
pc = read("src/server/utils/promptContext.ts")

BASE_PROMPT = extract_backtick(sh, "export const DEFAULT_SYSTEM_PROMPT =")
PROACTIVE = extract_backtick(sh, "const PROACTIVE_OFFERING_ADDITION =").strip()
REACTIVE = extract_backtick(sh, "const REACTIVE_ONLY_ADDITION =").strip()
RISK_PROMPT = extract_backtick(cd, "const RISK_ASSESSMENT_PROMPT =")
REDACTION_PROMPT = extract_backtick(red, "const prompt =").strip()
BATCH_ADDITION = extract_backtick(red, "const batchInstructions = prompt +").strip()
SESSION_NAME_SYS = extract_backtick(sn, 'role: "system",')
# normalize the TS indentation inside the session-name literal
SESSION_NAME_SYS = "\n".join(line.strip() for line in SESSION_NAME_SYS.splitlines()).strip()

# Modality presets: label, addition, phases
presets = []
preset_block = sh[sh.index("DEFAULT_MODALITY_PRESETS"):sh.index("/** Resolve the active modality preset")]
for m in re.finditer(r"label: '([^']*)',\s*\n\s*addition: `(.*?)`,\s*\n\s*phases: \[(.*?)\n\s*\],", preset_block, re.S):
    label, addition, phases_src = m.group(1), decode_js(m.group(2)).strip(), m.group(3)
    phases = [
        (pm.group(1), pm.group(2), decode_js(pm.group(3)))
        for pm in re.finditer(r"\{ at: ([0-9.]+), label: '([^']*)', guidance: '((?:[^'\\]|\\.)*)' \}", phases_src)
    ]
    presets.append((label, addition, phases))
assert len(presets) == 4, f"expected 4 modality presets, got {len(presets)}"

# Tool-guidance lines from promptContext.ts (in source order, with the tool
# condition that gates each line).
tool_lines = []
tg_block = pc[pc.index("export function buildToolGuidanceBlock"):pc.index("/** Rolling clinical case profile block")]
for m in re.finditer(r"if \((.*?)\) \{\s*\n\s*lines\.push\(\s*'((?:[^'\\]|\\.)*)'", tg_block):
    cond = re.findall(r"has\('([^']+)'\)", m.group(1))
    tool_lines.append((" or ".join(cond), decode_js(m.group(2))))
assert len(tool_lines) >= 14, f"expected >=14 tool guidance lines, got {len(tool_lines)}"

TOOL_GUIDANCE_HEADER = ("## Using your tools (important)\n"
                        "Your function tools show real interactive cards and forms on the participant's screen and "
                        "save information for their care team. When a request matches a tool, CALL the tool — "
                        "describing it verbally instead is a failure. Specifically:")

CHECKIN_BLOCK = """## Today's check-in (provided by the participant just now)
- Current mood: <mood>/10
- On their mind: "<topic>"
- What they want from today: "<goal>"
Open the conversation by gently acknowledging this — don't interrogate them about it, just show you've heard it."""

MEMORY_BLOCK = """## Returning participant (conversation #<N> — they consented to session memory)
Context from recent conversations, most recent first:
- <date> ("<headline>"): topics: <topics> | <mood trajectory> | what helped: <techniques> | possible follow-up: <follow-up>
Things they explicitly asked you to remember:
- <fact the participant asked the agent to remember>
Clinical case profile (built from all prior sessions):
- Presenting concerns: <list>
- Recurring themes: <list>
- Stressors: <list>
- Support system: <list>
- Coping repertoire, most helpful first: <technique (helpfulness)>
- Values: <list>
- Screener trend: <text>
Since their last session:
- <SCALE>: <score> (was <previous score> — up/down/unchanged)
- Recent mood points: <mood>/10 -> <mood>/10 -> ...
- Has an existing safety plan (warning signs on file: <list>).
- Last thought record's balanced thought: "<text>"
Between-session practice:
- Open practice from last time: <title> (assigned <date>); ...
- They completed: <title>
Ask warmly how the practice went — celebrate follow-through, never scold a skipped one.
Guidance from the participant's care team (private — never read this aloud or reference it explicitly):
"<note left by the participant's therapist>"
Prior risk history (a therapist has enabled sharing this — use it ONLY to check in gently, never lead with it or list it back):
- <date>: <severity> severity (<later resolved/unflagged | no recorded resolution>)
If the conversation heads toward distress, you may check in warmly and vaguely ("how have you been holding up since we last talked?") — do not mention dates, scores, or that this history was flagged for you.
Use this for warmth and continuity ("last time we talked about..."), and to build on techniques that helped. Do not recite it back verbatim or claim to remember more than this."""

def extract_concat(src, anchor, occurrence=0):
    """Extract a possibly-concatenated sequence of backtick fragments starting
    after `anchor` (fragments joined by `+`, ended by anything else)."""
    i = -1
    for _ in range(occurrence + 1):
        i = src.index(anchor, i + 1)
    pos = i + len(anchor)
    out = []
    frag_re = re.compile(r"\s*\+?\s*`((?:[^`\\]|\\.)*)`")
    while True:
        m = frag_re.match(src, pos)
        if not m:
            break
        out.append(m.group(1))
        pos = m.end()
    if not out:
        raise SystemExit(f"extract_concat found nothing after {anchor!r}")
    return decode_js("".join(out))


ci = read("src/server/services/crisisIntervention.service.ts")
ms = read("src/server/services/minorSafeguard.service.ts")

steer_fn = ci[ci.index("function steeringGuidance"):ci.index("export function shouldSteer")]
STEER_BASE = extract_concat(steer_fn, "const base =").replace("${riskScore}", "<risk score>")
STEER_HIGH_SUFFIX = extract_concat(steer_fn, "return base +")
STEER_OTHER_SUFFIX = extract_concat(steer_fn, "return base +", occurrence=1)
chat_steer_fn = ci[ci.index("export function buildChatSteeringGuidance"):ci.index("CHAT_SAFETY_PROTOCOL_GUIDANCE")]
STEER_CHAT = extract_concat(chat_steer_fn, "return (").replace("${riskScore}", "<risk score>")

SAFETY_PROTOCOL = extract_concat(ci, "const SAFETY_PROTOCOL_GUIDANCE =")
CHAT_SAFETY_PROTOCOL = extract_concat(ci, "export const CHAT_SAFETY_PROTOCOL_GUIDANCE =")
WIND_DOWN = extract_concat(ci, "const CRISIS_WIND_DOWN_GUIDANCE =")
REALTIME_MINOR = extract_concat(ms, "export const REALTIME_MINOR_GUIDANCE =")
MINOR_MESSAGE = extract_concat(ms, "export const MINOR_ELIGIBILITY_MESSAGE =")

for name, val in [("STEER_BASE", STEER_BASE), ("SAFETY_PROTOCOL", SAFETY_PROTOCOL),
                  ("CHAT_SAFETY_PROTOCOL", CHAT_SAFETY_PROTOCOL), ("WIND_DOWN", WIND_DOWN),
                  ("REALTIME_MINOR", REALTIME_MINOR), ("MINOR_MESSAGE", MINOR_MESSAGE)]:
    if len(val) < 100:
        raise SystemExit(f"suspiciously short extraction for {name}: {val!r}")

# ---------------------------------------------------------------------------
# Document content model: list of (kind, text) items.
# kinds: title, meta, h1, h2, h3, p, note, pre
# ---------------------------------------------------------------------------
doc_items = []
add = lambda kind, text: doc_items.append((kind, text))

add("title", "AI Support Agent — Model System Prompts (Verbatim)")
add("meta", "Study: Human + AI Therapy: Better For Everyone (Phase 2: Longitudinal Use)\n"
            "Prepared: September 4, 2026 (supersedes the September 2, 2026 version)\n"
            "Sources: Application codebase — src/server/utils/sessionHelpers.ts, src/server/utils/promptContext.ts, "
            "src/server/services/crisisIntervention.service.ts, src/server/services/crisisDetection.service.ts, "
            "src/server/services/minorSafeguard.service.ts, src/server/services/redaction.service.ts, "
            "src/server/services/sessionName.service.ts.")

add("p", "This document reproduces, verbatim, every instruction text the platform sends to its AI models. "
         "Part A covers the per-session instruction set assembled when a participant session starts. Part B covers "
         "system-side guidance the platform can inject into a live session (never visible to the participant). "
         "Part C covers auxiliary prompts used by backend safety and data-processing pipelines. Part D lists the "
         "models in use. Administrators can update the deployed prompts through an audited admin interface, which "
         "can hold separate base prompts for voice (realtime) and typed-chat sessions; the text below is the "
         "deployed default, and the research team will keep the IRB-approved version and the deployed version in "
         "sync. For dynamic blocks that are filled with per-participant data at runtime, the fixed template text is "
         "reproduced verbatim and runtime values are marked with <angle-bracket placeholders>.")

add("h1", "Part A. Per-session instruction set")
add("h2", "A.1 How the instruction set is assembled")
add("p", "When a session starts, the system prompt sent to the conversation model is the concatenation, in order, of: "
         "(1) the base system prompt (A.2); "
         "(2) one therapeutic-approach appendix, only if a modality research condition is active (A.3); "
         "(3) exactly one exercise-offering appendix — proactive or reactive — assigned per session as a research "
         "condition (A.4); "
         "(4) a language addition, only for non-English sessions (A.5); "
         "(5) for voice sessions only, a tool-usage guidance block (A.6) — the typed-chat pipeline has no "
         "client-side tools and omits this block; "
         "(6) for logged-in returning participants who opted in to session memory, a returning-participant context "
         "block (A.7); and "
         "(7) a pre-session check-in block, only if the participant filled in the optional check-in form (A.8).")
add("p", "Voice and typed-chat sessions share this same assembly and the same default base prompt, but the stored "
         "configuration can hold a separate base prompt for each channel, and item (5) applies to voice only.")

add("h2", "A.2 Base system prompt (verbatim)")
add("p", "The {{crisis_text}} placeholder is interpolated at runtime from the configured crisis contacts as "
         "\"<hotline> <phone>, text <text line>, or 911\" — with the deployed configuration: "
         "\"988 Suicide & Crisis Lifeline 988, text Text HOME to 741741, or 911\". If the crisis-contact "
         "configuration is disabled it falls back to \"988 (Suicide and Crisis Lifeline), or 911 for immediate "
         "danger\".")
add("pre", BASE_PROMPT)

add("h2", "A.3 Therapeutic-approach appendices (applied only when a modality research condition is active)")
add("p", "At most one of the following four appendices is added, per the active modality condition (the default is "
         "none). Each appendix also defines session-pacing guidance lines that the platform delivers to the model "
         "at fractional points through the session (for example, at 15%, 45%, 70%, and 85% of the session's "
         "duration); these steer conversation structure only and introduce no new capabilities. Both the appendix "
         "and its pacing lines are reproduced verbatim below.")
for label, addition, phases in presets:
    add("h3", label)
    add("pre", addition)
    add("note", "Session-pacing guidance (delivered at the listed fraction of session time):")
    add("pre", "\n".join(f"[at {int(float(at) * 100)}% — {plabel}] {guidance}" for at, plabel, guidance in phases))

add("h2", "A.4 Exercise-offering appendices (per-session research condition, exactly one of two)")
add("h3", "Proactive condition")
add("pre", PROACTIVE)
add("h3", "Reactive condition")
add("pre", REACTIVE)

add("h2", "A.5 Language addition")
add("p", "For sessions conducted in a language other than English, the stored language configuration can append a "
         "short instruction (for example, to conduct the conversation in that language). English sessions — the "
         "study default — receive no language addition.")

add("h2", "A.6 Tool-usage guidance block (voice sessions only)")
add("p", "Voice sessions expose client-side function tools (on-screen exercises, mood logging, session memory, "
         "safety-plan retrieval, and similar). This block teaches the model when to actually call them. The header "
         "is always included; each bulleted line is included only when the corresponding tool is enabled for the "
         "session (the gating tool names are shown in brackets). Typed-chat sessions have no client-side tools and "
         "do not receive this block.")
add("pre", TOOL_GUIDANCE_HEADER + "\n" + "\n".join(line for _, line in tool_lines))
add("note", "Gating tools, in order of the lines above: " + "; ".join(cond for cond, _ in tool_lines) + ".")

add("h2", "A.7 Returning-participant context block (opt-in session memory)")
add("p", "Included only for logged-in participants who have enabled session memory and have prior context. The "
         "fixed template text below is verbatim; runtime values are marked with <placeholders>. Each labeled "
         "sub-section (recent conversations, remembered facts, case profile, since-last-session signals, "
         "between-session practice, care-team guidance, prior risk history) appears only when the corresponding "
         "data exists. The prior-risk-history sub-section is included only when a therapist has explicitly enabled "
         "sharing it for that participant.")
add("pre", MEMORY_BLOCK)

add("h2", "A.8 Pre-session check-in block")
add("p", "Included only when the participant fills in the optional pre-session check-in (mood 1-10, what's on "
         "their mind, what they want from today). Lines appear only for the fields the participant provided.")
add("pre", CHECKIN_BLOCK)

add("h1", "Part B. System-side guidance injected during a live session")
add("p", "During a session, the platform's safety pipeline can inject additional system-role guidance to the model. "
         "These messages are never shown to the participant; each begins with an instruction telling the model "
         "never to mention or acknowledge them. The trigger conditions are stated with each text.")

add("h2", "B.1 Risk-adaptive de-escalation steering")
add("p", "Trigger: the crisis-detection pipeline (Part C.1) scores a participant message at 25/100 or higher. "
         "Injected at most once per three minutes per session, shared across the voice and chat channels. "
         "<risk score> is the numeric score.")
add("h3", "Voice (realtime) sessions — severity below high")
add("pre", STEER_BASE + STEER_OTHER_SUFFIX)
add("h3", "Voice (realtime) sessions — high severity")
add("pre", STEER_BASE + STEER_HIGH_SUFFIX)
add("h3", "Typed-chat sessions")
add("pre", STEER_CHAT)

add("h2", "B.2 High-severity safety protocol")
add("p", "Trigger: the crisis-detection pipeline flags a message at high severity. The research team's monitoring "
         "dashboard is alerted and the on-call researcher is paged at the same time. The voice version leans on "
         "client-side tools (an on-screen resource card and a guided safety-plan builder); the chat version "
         "delivers resources directly in the reply.")
add("h3", "Voice (realtime) sessions")
add("pre", SAFETY_PROTOCOL)
add("h3", "Typed-chat sessions")
add("pre", CHAT_SAFETY_PROTOCOL)

add("h2", "B.3 Crisis wind-down (human-initiated graceful end)")
add("p", "Trigger: a member of the research team monitoring a crisis session decides it should end gracefully. The "
         "guidance below is injected to the voice model, which surfaces crisis resources and closes the session "
         "warmly; if the session has not ended within a 75-second grace window the server ends it directly.")
add("pre", WIND_DOWN)

add("h2", "B.4 Minor-eligibility ending")
add("p", "Trigger: a two-stage safeguard (keyword screen plus a gpt-4o-mini structured confirmation) confirms the "
         "participant has disclosed being under 18. The session is then ended. In voice sessions the model is "
         "instructed to deliver the goodbye itself; in typed chat the platform sends the participant a fixed, "
         "server-authored goodbye message instead.")
add("h3", "Voice (realtime) sessions — guidance injected to the model")
add("pre", REALTIME_MINOR)
add("h3", "Typed-chat sessions — fixed message shown to the participant")
add("pre", MINOR_MESSAGE)

add("h1", "Part C. Auxiliary backend prompts")

add("h2", "C.1 Crisis risk-assessment prompt (gpt-4o-mini)")
add("p", "Runs when a keyword screen flags possible crisis language in a participant message, and additionally as "
         "a periodic sweep after every eight participant messages without an assessment (so unusual phrasing is "
         "not missed). System prompt, verbatim:")
add("pre", RISK_PROMPT)
add("note", "The accompanying user message is: 'Recent transcript:' followed by the recent conversation, then "
            "'Latest participant message (the one that tripped the screen):' followed by the flagged message in "
            "quotes.")

add("h2", "C.2 Redaction prompt (gpt-5)")
add("p", "Every stored conversation is passed through this redaction prompt in two passes (the second pass re-checks "
         "the output of the first) before researchers can view it. Verbatim:")
add("pre", REDACTION_PROMPT)
add("note", "For whole-session batches, the following is appended to the prompt above, verbatim:")
add("pre", BATCH_ADDITION)

add("h2", "C.3 Session-name prompt (gpt-4o-mini)")
add("p", "Generates a short descriptive title for each session from the redacted transcript only. System prompt, "
         "verbatim:")
add("pre", SESSION_NAME_SYS)
add("note", "The accompanying user message is: 'Summarize this therapy session in 3-5 words:' followed by the "
            "redacted transcript (truncated to 3,000 characters).")

add("h1", "Part D. Models in use")
add("table", [
    ["Function", "Model"],
    ["Voice conversation (realtime)", "gpt-realtime-2.1-mini"],
    ["Voice transcription", "gpt-4o-mini-transcribe"],
    ["Typed chat conversation", "gpt-5.2"],
    ["Transcript redaction (dual-pass)", "gpt-5"],
    ["Auxiliary safety and analysis (crisis risk assessment, minor-eligibility confirmation, session naming)", "gpt-4o-mini"],
    ["Text embeddings (memory retrieval)", "text-embedding-3-small"],
])

# ---------------------------------------------------------------------------
# Emit DOCX
# ---------------------------------------------------------------------------
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

for kind, payload in doc_items:
    if kind == "title":
        p = doc.add_heading(payload, level=0)
    elif kind == "meta":
        for line in payload.split("\n"):
            p = doc.add_paragraph()
            key, _, rest = line.partition(": ")
            r = p.add_run(key + ": ")
            r.bold = True
            p.add_run(rest)
            p.paragraph_format.space_after = Pt(2)
    elif kind == "h1":
        doc.add_heading(payload, level=1)
    elif kind == "h2":
        doc.add_heading(payload, level=2)
    elif kind == "h3":
        doc.add_heading(payload, level=3)
    elif kind == "p":
        doc.add_paragraph(payload)
    elif kind == "note":
        p = doc.add_paragraph()
        r = p.add_run(payload)
        r.italic = True
        r.font.size = Pt(9.5)
    elif kind == "pre":
        for line in payload.split("\n"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.25)
            r = p.add_run(line if line else " ")
            r.font.name = "Consolas"
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
    elif kind == "table":
        t = doc.add_table(rows=0, cols=2)
        t.style = "Table Grid"
        for i, row in enumerate(payload):
            cells = t.add_row().cells
            for j, val in enumerate(row):
                cells[j].text = val
                if i == 0:
                    for r in cells[j].paragraphs[0].runs:
                        r.bold = True

doc.save(OUT_DOCX)
print(f"wrote {OUT_DOCX}")

# ---------------------------------------------------------------------------
# Emit HTML
# ---------------------------------------------------------------------------
esc = html_mod.escape
parts = ['<html><head><meta charset="utf-8"><title>AI Model System Prompts</title>',
         "<style>body{font-family:Georgia,serif;max-width:52em;margin:2em auto;padding:0 1em;color:#222;}"
         "pre{background:#f6f6f4;border:1px solid #ddd;padding:0.8em;white-space:pre-wrap;"
         "font-size:0.82em;line-height:1.45;}"
         ".note{font-style:italic;font-size:0.9em;color:#444;}"
         "table{border-collapse:collapse;}td,th{border:1px solid #999;padding:0.35em 0.7em;text-align:left;}"
         "</style></head><body>"]
for kind, payload in doc_items:
    if kind == "title":
        parts.append(f"<h1>{esc(payload)}</h1>")
    elif kind == "meta":
        lines = []
        for line in payload.split("\n"):
            key, _, rest = line.partition(": ")
            lines.append(f"<b>{esc(key)}:</b> {esc(rest)}")
        parts.append("<p>" + "<br>\n".join(lines) + "</p>")
    elif kind == "h1":
        parts.append(f"<h2>{esc(payload)}</h2>")
    elif kind == "h2":
        parts.append(f"<h3>{esc(payload)}</h3>")
    elif kind == "h3":
        parts.append(f"<h4>{esc(payload)}</h4>")
    elif kind == "p":
        parts.append(f"<p>{esc(payload)}</p>")
    elif kind == "note":
        parts.append(f'<p class="note">{esc(payload)}</p>')
    elif kind == "pre":
        parts.append(f"<pre>{esc(payload)}</pre>")
    elif kind == "table":
        rows = []
        for i, row in enumerate(payload):
            tag = "th" if i == 0 else "td"
            rows.append("<tr>" + "".join(f"<{tag}>{esc(v)}</{tag}>" for v in row) + "</tr>")
        parts.append("<table>" + "\n".join(rows) + "</table>")
parts.append("</body></html>")
with open(OUT_HTML, "w", encoding="utf-8") as f:
    f.write("\n".join(parts) + "\n")
print(f"wrote {OUT_HTML}")
